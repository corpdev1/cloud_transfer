"""
Generates the drivetocloud end-user guide as a formatted .docx file.
Run:  python make_docs.py
Output: drivetocloud_guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6C)   # headings
MID_BLUE   = RGBColor(0x20, 0x5C, 0xA8)   # sub-headings / links
ACCENT     = RGBColor(0x16, 0x7A, 0xCE)   # step numbers / callout border
LIGHT_GREY = RGBColor(0xF2, 0xF4, 0xF8)   # code block background
GREEN      = RGBColor(0x1A, 0x7A, 0x3C)   # success / tip label
ORANGE     = RGBColor(0xC0, 0x5C, 0x00)   # warning label
RED        = RGBColor(0xB0, 0x1C, 0x1C)   # important / caution
TEXT       = RGBColor(0x1C, 0x1C, 0x1C)   # body text
DIM        = RGBColor(0x55, 0x55, 0x55)   # secondary text

# ── Style helpers ─────────────────────────────────────────────────────────────

def _set_font(run, name="Calibri", size=11, bold=False, italic=False,
              color=TEXT, underline=False):
    run.font.name      = name
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color
    run.font.underline = underline


def _para_space(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _shade_cell(cell, hex_color="F2F4F8"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color="167ACE", size="12"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("left",):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


# ── Building blocks ───────────────────────────────────────────────────────────

def heading1(text):
    p = doc.add_paragraph()
    _para_space(p, before=18, after=6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_font(run, size=20, bold=True, color=DARK_BLUE)
    return p


def heading2(text):
    p = doc.add_paragraph()
    _para_space(p, before=14, after=4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_font(run, size=14, bold=True, color=DARK_BLUE)
    return p


def heading3(text):
    p = doc.add_paragraph()
    _para_space(p, before=10, after=3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_font(run, size=12, bold=True, color=MID_BLUE)
    return p


def body(text, bold_parts=None):
    """
    body("Hello **world** today")  — bold_parts auto-parsed from **...**
    or body("plain text")
    """
    p = doc.add_paragraph()
    _para_space(p, after=5)
    p.paragraph_format.left_indent = Cm(0)
    _write_inline(p, text)
    return p


def _write_inline(para, text):
    """Parse **bold** and `code` markers and add runs accordingly."""
    import re
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for tok in tokens:
        if tok.startswith("**") and tok.endswith("**"):
            run = para.add_run(tok[2:-2])
            _set_font(run, bold=True)
        elif tok.startswith("`") and tok.endswith("`"):
            run = para.add_run(tok[1:-1])
            _set_font(run, name="Courier New", size=10, color=MID_BLUE)
        else:
            run = para.add_run(tok)
            _set_font(run)


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    _para_space(p, after=3)
    indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.left_indent  = indent
    p.paragraph_format.first_line_indent = Cm(-0.4)
    _write_inline(p, text)
    return p


def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    _para_space(p, after=4)
    indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.left_indent  = indent
    p.paragraph_format.first_line_indent = Cm(-0.5)
    _write_inline(p, text)
    return p


def code_block(lines):
    """Render a shaded monospace block."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "EAEEF5")
    cell.width = Inches(5.8)
    # clear default paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for i, line in enumerate(lines):
        cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(1) if i > 0 else Pt(5)
        cp.paragraph_format.space_after  = Pt(1) if i < len(lines)-1 else Pt(5)
        cp.paragraph_format.left_indent  = Cm(0.3)
        run = cp.add_run(line)
        _set_font(run, name="Courier New", size=10, color=RGBColor(0x10, 0x3A, 0x6E))
    doc.add_paragraph()   # spacing after table


def callout(label, text, label_color=GREEN, border_hex="1A7A3C"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "F0FAF4" if label_color == GREEN else
                       "FFF8ED" if label_color == ORANGE else "FDF0F0")
    _set_cell_border(cell, color=border_hex, size="18")
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    lp = cell.add_paragraph()
    lp.paragraph_format.space_before = Pt(5)
    lp.paragraph_format.space_after  = Pt(2)
    lp.paragraph_format.left_indent  = Cm(0.4)
    lr = lp.add_run(label)
    _set_font(lr, bold=True, color=label_color, size=10)

    tp = cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(5)
    tp.paragraph_format.left_indent  = Cm(0.4)
    _write_inline(tp, text)
    doc.add_paragraph()


def step_header(n, title):
    p = doc.add_paragraph()
    _para_space(p, before=12, after=4)
    p.paragraph_format.keep_with_next = True
    num_run = p.add_run(f"Step {n}  ")
    _set_font(num_run, size=13, bold=True, color=ACCENT)
    title_run = p.add_run(title)
    _set_font(title_run, size=13, bold=True, color=DARK_BLUE)


def divider():
    p = doc.add_paragraph()
    _para_space(p, before=4, after=4)
    run = p.add_run("─" * 72)
    _set_font(run, size=8, color=RGBColor(0xCC, 0xCC, 0xCC))


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════


# ── Section 1 ─────────────────────────────────────────────────────────────────
heading1("1 — What this tool does")
body("**drivetocloud** copies files from Google Drive — or a folder on your computer — to a storage server. It runs from a terminal (command line) but the setup wizard guides you through configuration in plain steps.")
body("Once set up, you run three simple commands:")
numbered("**Scan** — the tool looks at your files and makes a list of what needs to be transferred.")
numbered("**Transfer** — it uploads each file. You can stop at any time and resume later; it picks up exactly where it left off.")
numbered("**Status** — shows you a live count of how many files are done, pending, or failed.")
callout("TIP", "You do not need to understand programming or networking to use this tool. Follow the wizard and it will configure everything for you.", GREEN, "1A7A3C")

page_break()

# ── Section 2 ─────────────────────────────────────────────────────────────────
heading1("2 — Before you begin — what you need")
heading2("On your computer")
bullet("**Python 3.10 or later** — the programming language the tool runs on")
bullet("**Git** — used to download the tool from GitHub")
bullet("**A terminal app** — Terminal on Mac, Command Prompt or PowerShell on Windows")
bullet("**Internet connection**")

body("**How to check if Python and Git are already installed** — open your terminal and type:")
code_block(["python --version", "git --version"])
body("Both should print a version number. If either says 'command not found', download and install them first:")
bullet("Python: **python.org/downloads** — download the latest version and run the installer")
bullet("Git: **git-scm.com/downloads** — download and run the installer (use all default options)")

heading2("From Google (only if transferring from Google Drive)")
bullet("An **OAuth client file** (for personal Drive) **or** a **service account key file** (for organisation Drive) — your administrator will provide this or tell you how to get it")

heading2("From your administrator")
bullet("The file called **`operator.env`** — this contains the storage server details and must be placed inside the tool folder before you run the wizard")

callout("IMPORTANT",
        "Do not open or edit operator.env. It contains connection details for the storage server. Just place it in the drivetocloud folder as given to you.",
        ORANGE, "C05C00")

page_break()

# ── Section 3 ─────────────────────────────────────────────────────────────────
heading1("3 — Install the tool on your computer")

step_header(1, "Open your terminal")
bullet("**Mac:** Press Cmd + Space, type **Terminal**, press Enter")
bullet("**Windows:** Press the Windows key, type **cmd**, press Enter")

step_header(2, "Download the tool from GitHub")
body("Copy and paste this command exactly — it will create a folder called `cloud_transfer` wherever your terminal is currently open:")
code_block(["git clone https://github.com/corpdev1/cloud_transfer.git"])
body("You should see output like:")
code_block([
    "Cloning into 'cloud_transfer'...",
    "remote: Enumerating objects: 11, done.",
    "Resolving deltas: 100% — done.",
])
callout("TIP", "By default this creates the folder in your home directory. To put it somewhere specific, open Finder/File Explorer, navigate to that folder, then open the terminal there before running the command above.", GREEN, "1A7A3C")

step_header(3, "Enter the tool folder")
body("Navigate into the folder that was just created:")
code_block(["cd cloud_transfer"])

step_header(4, "Install required libraries")
body("This only needs to be done once. Copy and paste this exactly:")
code_block(["pip install -r requirements.txt"])
body("You should see a list of packages being installed. Wait until it finishes and you see the `$` prompt again.")
callout("TIP", "If you see 'pip: command not found', try `pip3 install -r requirements.txt` instead.", GREEN, "1A7A3C")

step_header(5, "Place the operator.env file")
body("Copy the **`operator.env`** file your administrator gave you into the `cloud_transfer` folder. It should sit alongside `main.py`, `wizard.py`, etc.")
callout("NOTE", "If operator.env is missing, the wizard will still run but will ask you for storage server details — contact your administrator if you do not have these.", ORANGE, "C05C00")

page_break()

# ── Section 4 ─────────────────────────────────────────────────────────────────
heading1("4 — Run the Setup Wizard")
body("The wizard sets up your personal configuration. It takes about 2 minutes. Run it once before your first transfer.")
code_block(["python wizard.py"])
body("The wizard will ask you two questions:")
numbered("**Where are your files?** — Google Drive (personal), Google Drive (organisation), or a local folder")
numbered("**Where do you want to upload them?** — Your administrator's pre-configured storage, your own SFTP server, or your own S3 storage")
body("Then it saves a file called **`.env`** with your settings. You do not need to edit this file manually.")

divider()
heading2("4a — Uploading from Google Drive (personal account)")

step_header(1, "Choose source")
body("When the wizard asks 'Where are your files?', select:")
code_block(["1. Google Drive — personal account  (sign in with browser)"])

step_header(2, "Get your OAuth client file")
body("If your administrator has not already provided this file, you need to create one:")
numbered("Go to **console.cloud.google.com** and sign in with your Google account")
numbered("Click **Select a project** at the top, then **New Project** — give it any name")
numbered("In the left menu, go to **APIs & Services → Library**")
numbered("Search for **Google Drive API** and click **Enable**")
numbered("Go to **APIs & Services → Credentials**")
numbered("Click **+ Create Credentials → OAuth 2.0 Client ID**")
numbered("Choose **Desktop app** as the application type, give it a name, click **Create**")
numbered("Click **Download JSON** — save this file into the `credentials/` folder inside drivetocloud")
numbered("Rename it to `oauth_client.json` (optional but keeps things tidy)")

step_header(3, "Enter the file path in the wizard")
body("When asked, type:")
code_block(["credentials/oauth_client.json"])

step_header(4, "Sign in to Google in the browser")
body("When you later run `python main.py enumerate` or `python main.py list-drives` for the first time, a browser window will open automatically. Sign in with your Google account and click **Allow**. This only happens once.")

callout("TIP", "If the browser does not open automatically, the terminal will show a URL — copy and paste it into your browser manually.", GREEN, "1A7A3C")

divider()
heading2("4b — Uploading from Google Drive (organisation / workspace)")

body("This option lets the tool access files across your entire organisation without needing individual logins. It requires a **service account key** — your IT administrator will usually provide this.")

step_header(1, "Choose source")
body("When the wizard asks 'Where are your files?', select:")
code_block(["2. Google Drive — organization / workspace  (service account key)"])

step_header(2, "Get your service account key file")
body("If your administrator has not already provided this file:")
numbered("Go to **console.cloud.google.com** and sign in")
numbered("Go to **IAM & Admin → Service Accounts**")
numbered("Select or create a service account")
numbered("Click on the account, go to the **Keys** tab")
numbered("Click **Add Key → Create new key → JSON**")
numbered("Save the downloaded file into the `credentials/` folder inside drivetocloud")

step_header(3, "Enter the file path in the wizard")
code_block(["credentials/sa_key.json"])

step_header(4, "Impersonation email (optional)")
body("If the wizard asks for an 'email to impersonate', enter the email of a workspace admin (e.g. `admin@yourcompany.com`). This is needed for workspace-wide access. Press Enter to skip if unsure.")

callout("IMPORTANT",
        "The service account must have domain-wide delegation enabled by your Google Workspace admin before it can impersonate users. If you get a 'permission denied' error, contact your IT administrator.",
        ORANGE, "C05C00")

divider()
heading2("4c — Uploading from a folder on your computer")

step_header(1, "Choose source")
body("When the wizard asks 'Where are your files?', select:")
code_block(["3. Local folder on this computer"])

step_header(2, "Enter the folder path")
body("Type the full path to the folder you want to upload. Examples:")
code_block([
    "Mac / Linux:   /Users/yourname/Documents/MyFiles",
    "Windows:       C:\\Users\\YourName\\Documents\\MyFiles",
])
callout("TIP", "On Mac or Linux, you can drag the folder from Finder/File Manager into the terminal window — it will paste the path automatically.", GREEN, "1A7A3C")

step_header(3, "Choose your folder name")
body("The wizard will ask for a folder name — this is the name your files will be stored under on the storage server. Choose something descriptive like your name or the project name (e.g. `rohit_files` or `q1_reports`).")

page_break()

# ── Section 5 ─────────────────────────────────────────────────────────────────
heading1("5 — Scan your files")
body("Before transferring, the tool needs to build a list of all your files. This is called **enumeration**. It does not upload anything — it just catalogues what needs to be transferred.")

heading2("For Google Drive files")
body("First, see a list of your available drives:")
code_block(["python main.py list-drives"])
body("Then scan a specific folder by pasting its URL:")
code_block(['python main.py enumerate --folder "https://drive.google.com/drive/folders/FOLDER_ID"'])
body("Or scan a specific shared drive by name:")
code_block(['python main.py enumerate --drives "Marketing Drive"'])
body("Or scan everything:")
code_block(["python main.py enumerate"])

heading2("For a local folder")
code_block(['python main.py enumerate-local "/path/to/your/folder"'])

body("After scanning, check what was found:")
code_block(["python main.py status"])
body("You will see a table like this:")
code_block([
    "Status           Files            Size",
    "─────────────────────────────────────────",
    "done                 0            0.0 B",
    "pending          3,214           42.7 GB",
    "failed               0            0.0 B",
    "─────────────────────────────────────────",
    "TOTAL            3,214           42.7 GB",
])

callout("TIP", "You can run enumerate multiple times — it will not create duplicates. It is safe to re-scan if you added more files.", GREEN, "1A7A3C")

page_break()

# ── Section 6 ─────────────────────────────────────────────────────────────────
heading1("6 — Start the transfer")
body("Once scanning is complete, start uploading:")
code_block(["python main.py transfer"])
body("You will see live progress in the terminal, updated every few files:")
code_block([
    "Progress: 250/3,214 (7.8%) | done=250 failed=0 transferred=3.2 GB speed=18.4 MB/s eta=0h38m",
])
body("The transfer is **safe to stop at any time** — press `Ctrl + C` to pause. When you run `python main.py transfer` again, it will resume from where it stopped.")

callout("TIP", "For faster transfers, increase PARALLEL_WORKERS in your .env file (e.g. set it to 8 or 16). Higher values use more memory and bandwidth.", GREEN, "1A7A3C")

callout("NOTE", "If you are transferring from Google Drive for the first time, a browser window will open asking you to sign in to Google. This is normal — sign in and click Allow.", ORANGE, "C05C00")

page_break()

# ── Section 7 ─────────────────────────────────────────────────────────────────
heading1("7 — Check progress")
body("You can check progress at any time — even while the transfer is running in another window:")
code_block(["python main.py status"])
body("The status table shows:")
bullet("**done** — files successfully uploaded")
bullet("**pending** — files waiting to be uploaded")
bullet("**in_progress** — files currently being uploaded")
bullet("**failed** — files that could not be uploaded (see Section 8)")

body("When the transfer finishes, the terminal will print a summary:")
code_block([
    "Transfer complete — done: 3,214 (42.7 GB), failed: 0",
])

page_break()

# ── Section 8 ─────────────────────────────────────────────────────────────────
heading1("8 — If something goes wrong")

heading2("Some files failed to upload")
body("Check which files failed and why:")
code_block(["python main.py show-failed"])
body("Then reset them and try again:")
code_block(["python main.py retry-failed", "python main.py transfer"])

heading2("The transfer stopped unexpectedly")
body("Just run transfer again — it automatically resumes from where it stopped:")
code_block(["python main.py transfer"])

heading2("You want to start completely fresh")
body("Delete the state database file (e.g. `transfer_state.db`) and re-run enumerate:")
code_block(["python main.py enumerate", "python main.py transfer"])
callout("IMPORTANT", "Deleting the state DB means the tool loses track of what was already uploaded. Only do this if you want to start over completely.", ORANGE, "C05C00")

heading2("Google Drive login expired")
body("If you see a message about an expired or invalid token, delete the cached login and re-authenticate:")
code_block(["rm credentials/token.json", "python main.py list-drives"])
body("A browser window will open again for you to sign in.")

page_break()

# ── Section 9 ─────────────────────────────────────────────────────────────────
heading1("9 — Common error messages explained")

def error_entry(error_text, explanation, fix):
    p = doc.add_paragraph()
    _para_space(p, before=8, after=2)
    r = p.add_run(error_text)
    _set_font(r, name="Courier New", size=10, color=RED)
    body(f"**What it means:** {explanation}")
    body(f"**Fix:** {fix}")
    doc.add_paragraph()

error_entry(
    "OAuth client file not found: credentials/oauth_client.json",
    "The Google sign-in file is missing.",
    "Download it from Google Cloud Console (see Section 4a) and save it to the credentials/ folder."
)
error_entry(
    "Service account key file not found",
    "The organisation login file is missing.",
    "Get the service account JSON key from your IT administrator and save it to the credentials/ folder."
)
error_entry(
    "insufficientPermissions",
    "The logged-in account does not have permission to access the requested Drive or folder.",
    "Make sure you are signed in with the correct Google account, or ask your administrator to share the Drive with you."
)
error_entry(
    "Cannot create remote directory … no space",
    "The storage server is full.",
    "Contact your administrator — the storage box needs more space."
)
error_entry(
    "Token has been expired or revoked",
    "Your Google login session has expired.",
    "Delete credentials/token.json and run python main.py list-drives to sign in again."
)
error_entry(
    "Missing env var: SFTP_HOST",
    "The storage server address is not configured.",
    "Make sure operator.env is in the drivetocloud folder. Ask your administrator for this file."
)
error_entry(
    "PARALLEL_WORKERS must be >= 1",
    "An invalid value was set in .env.",
    "Open .env in a text editor and make sure PARALLEL_WORKERS is a number like 4 or 8."
)
error_entry(
    "Drive name 'X' matches multiple drives",
    "Two shared drives have the same name (differing only in upper/lower case).",
    "Run python main.py list-drives, copy the drive ID (the long string), and use that instead of the name."
)

# ── Final page ────────────────────────────────────────────────────────────────
page_break()
heading1("Quick Reference")
body("All commands at a glance:")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
_shade_cell(hdr[0], "1A376C")
_shade_cell(hdr[1], "1A376C")
for cell, text in [(hdr[0], "Command"), (hdr[1], "What it does")]:
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _set_font(run, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=11)

rows = [
    ("python wizard.py",                         "First-time setup — configure source and destination"),
    ("python main.py list-drives",               "Show all Google Drive folders you can access"),
    ("python main.py enumerate",                 "Scan all drives and queue files for transfer"),
    ('python main.py enumerate --folder "URL"',  "Scan a specific Google Drive folder by URL"),
    ('python main.py enumerate --drives "Name"', "Scan a specific shared drive by name"),
    ('python main.py enumerate-local "/path"',   "Scan a local folder and queue for transfer"),
    ("python main.py status",                    "Show transfer progress (files done / pending / failed)"),
    ("python main.py transfer",                  "Start or resume the transfer"),
    ("python main.py retry-failed",              "Reset failed files so they are tried again"),
    ("python main.py show-failed",               "Show which files failed and the error message"),
]
for cmd, desc in rows:
    row = tbl.add_row().cells
    p0 = row[0].paragraphs[0]
    run0 = p0.add_run(cmd)
    _set_font(run0, name="Courier New", size=9, color=MID_BLUE)
    p1 = row[1].paragraphs[0]
    run1 = p1.add_run(desc)
    _set_font(run1, size=10)

doc.add_paragraph()
callout("NEED HELP?",
        "Contact your administrator or re-run python wizard.py to reset your configuration.",
        MID_BLUE, "205CA8")

# ── Save ──────────────────────────────────────────────────────────────────────
out = "drivetocloud_guide.docx"
doc.save(out)
print(f"Saved: {out}")
